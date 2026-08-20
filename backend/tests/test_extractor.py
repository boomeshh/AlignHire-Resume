import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch
import docx
from backend.app.extractor import extract_text

def test_extract_txt(tmp_path):
    # Create temporary text file
    txt_file = tmp_path / "resume.txt"
    content = "John Doe\njohn@example.com"
    txt_file.write_text(content, encoding="utf-8")

    result = extract_text(str(txt_file))

    assert isinstance(result, str)
    assert "John Doe" in result
    assert "john@example.com" in result

def test_extract_txt_utf8(tmp_path):
    # Test UTF-8 handling specifically with special unicode characters
    txt_file = tmp_path / "resume_utf8.txt"
    content = "John Doe • Python Developer • Résumé"
    txt_file.write_text(content, encoding="utf-8")
    
    result = extract_text(str(txt_file))
    assert "Résumé" in result
    assert "•" in result

def test_unsupported_file():
    with pytest.raises(ValueError):
        extract_text("resume.xyz")

def test_missing_file():
    with pytest.raises(FileNotFoundError):
        extract_text("missing.pdf")

def test_extract_docx_paragraphs_and_tables(tmp_path):
    doc_path = tmp_path / "resume.docx"
    doc = docx.Document()
    doc.add_paragraph("This is normal paragraph text.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill Name"
    table.cell(0, 1).text = "Experience Level"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Expert"
    doc.save(str(doc_path))
    
    result = extract_text(str(doc_path))
    assert "This is normal paragraph text." in result
    assert "Skill Name" in result
    assert "Experience Level" in result
    assert "Python" in result
    assert "Expert" in result

def test_extract_pdf_empty_and_throwing_pages(tmp_path):
    # PDF file must exist on disk so Path.exists() check passes
    pdf_path = tmp_path / "resume.pdf"
    pdf_path.write_text("dummy pdf", encoding="utf-8")
    
    # Mock pypdf.PdfReader to simulate normal, empty (None), and corrupt/crashing pages
    mock_reader = MagicMock()
    
    mock_page_normal = MagicMock()
    mock_page_normal.extract_text.return_value = "Normal page text."
    
    mock_page_none = MagicMock()
    mock_page_none.extract_text.return_value = None
    
    mock_page_crash = MagicMock()
    mock_page_crash.extract_text.side_effect = Exception("Page extraction crashed")
    
    mock_reader.pages = [mock_page_normal, mock_page_none, mock_page_crash]
    
    with patch("pypdf.PdfReader", return_value=mock_reader):
        result = extract_text(str(pdf_path))
        
    # Result should successfully contain the normal page text without crashing
    assert "Normal page text." in result
    # It must not include None or other failures
    assert len(result.strip()) > 0
