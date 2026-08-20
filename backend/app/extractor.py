import os
from pathlib import Path
import pypdf
import docx

def extract_text(file_path: str) -> str:
    """
    Extracts raw text from a given PDF, DOCX, or TXT file.
    Preserves line breaks and handles errors.
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    
    # 1. Check file extension first (strictly .pdf, .docx, .txt)
    if ext not in (".pdf", ".docx", ".txt"):
        raise ValueError(f"Unsupported file format: {ext}")
        
    # 2. Check if the file exists
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if not path.is_file():
        raise FileNotFoundError(f"Path is not a file: {file_path}")
        
    # 3. Extract based on extension
    if ext == ".pdf":
        text = _extract_from_pdf(path)
    elif ext == ".docx":
        text = _extract_from_docx(path)
    elif ext == ".txt":
        text = _extract_from_txt(path)
        
    # 3. Do not silently return empty output
    if not text.strip():
        raise ValueError(f"Extracted text is empty or invalid from file: {file_path}")
        
    return text

def _extract_from_pdf(path: Path) -> str:
    text_parts = []
    try:
        reader = pypdf.PdfReader(str(path))
        for page in reader.pages:
            try:
                # Safe handling for page.extract_text() returning None
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
            except Exception:
                # Empty PDF pages must not crash
                continue
    except Exception as e:
        raise ValueError(f"Failed to parse PDF file: {e}")
        
    return "\n".join(text_parts)

def _extract_from_docx(path: Path) -> str:
    try:
        doc = docx.Document(str(path))
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {e}")

def _extract_from_txt(path: Path) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        raise ValueError(f"Failed to parse TXT file: {e}")
